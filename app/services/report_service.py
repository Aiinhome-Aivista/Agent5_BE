"""
Report Service — generates the Weekly Optimization Report (Word).
"""
from __future__ import annotations

import io
import logging
from datetime import datetime, timedelta
from typing import Any, Dict

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy import func, desc
from sqlalchemy.orm import Session

from app.models.recommendation import Recommendation
from app.models.action_history import ActionHistory
from app.models.telemetry import CostSnapshot, AnomalyEvent

logger = logging.getLogger(__name__)


def generate_weekly_report(db: Session) -> io.BytesIO:
    """Generate a Word doc report and return as in-memory bytes."""
    end = datetime.utcnow()
    start = end - timedelta(days=7)
    prev_start = start - timedelta(days=7)

    # Aggregate stats
    total_spend = (
        db.query(func.sum(CostSnapshot.cost_usd))
        .filter(CostSnapshot.usage_date >= start)
        .scalar() or 0
    )
    prev_spend = (
        db.query(func.sum(CostSnapshot.cost_usd))
        .filter(CostSnapshot.usage_date >= prev_start,
                CostSnapshot.usage_date < start)
        .scalar() or 0
    )
    realized_savings = (
        db.query(func.sum(ActionHistory.realized_savings_usd))
        .filter(ActionHistory.completed_at >= start,
                ActionHistory.status == "succeeded")
        .scalar() or 0
    )
    executed = (
        db.query(ActionHistory)
        .filter(ActionHistory.completed_at >= start,
                ActionHistory.status.in_(["succeeded", "dry_run"]))
        .count()
    )
    pending_recs = (
        db.query(Recommendation)
        .filter(Recommendation.status == "pending")
        .order_by(Recommendation.estimated_monthly_savings_usd.desc())
        .limit(10)
        .all()
    )
    anomalies = (
        db.query(AnomalyEvent)
        .filter(AnomalyEvent.detected_at >= start)
        .order_by(AnomalyEvent.detected_at.desc())
        .limit(10)
        .all()
    )

    # Build doc
    doc = Document()

    # Title
    title = doc.add_heading("Weekly Optimization Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        f"Platform Performance & Cost Optimization Agent\n"
        f"Period: {start.date()} — {end.date()}"
    )
    sub_run.italic = True

    # Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    change_pct = ((total_spend - prev_spend) / prev_spend * 100) if prev_spend else 0
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = "Light Grid Accent 1"
    tbl.rows[0].cells[0].text = "Total spend (this week)"
    tbl.rows[0].cells[1].text = f"${total_spend:,.2f}"
    tbl.rows[1].cells[0].text = "Prior week spend"
    tbl.rows[1].cells[1].text = f"${prev_spend:,.2f}"
    tbl.rows[2].cells[0].text = "Week-over-week change"
    tbl.rows[2].cells[1].text = f"{change_pct:+.1f}%"
    tbl.rows[3].cells[0].text = "Realized savings (executed actions)"
    tbl.rows[3].cells[1].text = f"${realized_savings:,.2f}"

    # Anomalies
    doc.add_heading("2. Top Anomalies Detected", level=1)
    if not anomalies:
        doc.add_paragraph("No anomalies detected this week.")
    else:
        for a in anomalies:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f"[{a.severity.upper()}] {a.title}\n")
            r.bold = True
            p.add_run(f"  Detected: {a.detected_at.date()} | {a.provider} | "
                      f"Type: {a.anomaly_type}\n")
            if a.description:
                p.add_run(f"  {a.description}")

    # Pending recommendations
    doc.add_heading("3. Pending Recommendations", level=1)
    if not pending_recs:
        doc.add_paragraph("No pending recommendations.")
    else:
        recs_tbl = doc.add_table(rows=1, cols=4)
        recs_tbl.style = "Light Grid Accent 1"
        hdr = recs_tbl.rows[0].cells
        hdr[0].text = "Title"
        hdr[1].text = "Risk"
        hdr[2].text = "Est. Savings/mo"
        hdr[3].text = "Provider"
        for r in pending_recs:
            row = recs_tbl.add_row().cells
            row[0].text = r.title[:80]
            row[1].text = r.risk_class.upper()
            row[2].text = f"${r.estimated_monthly_savings_usd:.2f}"
            row[3].text = r.provider

    # Executed actions
    doc.add_heading("4. Actions Executed", level=1)
    actions = (
        db.query(ActionHistory)
        .filter(ActionHistory.completed_at >= start)
        .order_by(ActionHistory.completed_at.desc())
        .limit(20)
        .all()
    )
    if not actions:
        doc.add_paragraph("No actions executed this week.")
    else:
        for a in actions:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{a.status.upper()}] {a.action_type} on {a.resource_id} "
                      f"by {a.actor} at {a.started_at}")

    # Forward look
    doc.add_heading("5. Forward-Looking Forecast", level=1)
    next_week_proj = total_spend  # naive forecast
    pending_savings_pot = sum(r.estimated_monthly_savings_usd for r in pending_recs) / 4.0
    doc.add_paragraph(
        f"Projected spend next week: ${next_week_proj:,.2f} (flat baseline). "
        f"Approving all pending recommendations would unlock approximately "
        f"${pending_savings_pot:,.2f} of weekly savings ($"
        f"{sum(r.estimated_monthly_savings_usd for r in pending_recs):.2f}/mo)."
    )

    # Footer
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(f"\nGenerated: {datetime.utcnow().isoformat()}Z")
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
